#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re

# Read the HTML file
with open("index.html", "r", encoding="utf-8") as f:
    html_content = f.read()

soup = BeautifulSoup(html_content, "html.parser")


def create_document(language="he"):
    doc = Document()

    # Set document margins for printing
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Title
    title = doc.add_heading("Thai Garden", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(46, 125, 50)
    title_run.bold = True

    # Welcome message
    if language == "he":
        welcome_text = soup.find("div", class_="welcome-message")
        if welcome_text:
            welcome_para = doc.add_paragraph()
            welcome_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            welcome_run = welcome_para.add_run("ברוכים הבאים ל־Thai Garden 🌿\n")
            welcome_run.bold = True
            welcome_run.font.size = Pt(14)
            welcome_run.font.color.rgb = RGBColor(46, 125, 50)

            # Get welcome message text
            welcome_p = welcome_text.find_all("p")
            for p in welcome_p:
                text = p.get_text(strip=True)
                if text:
                    para = doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para_run = para.runs[0]
                    para_run.font.size = Pt(11)
            doc.add_paragraph()  # Empty line

    # Find food tab content
    if language == "he":
        food_content = soup.find("div", id="food")
        content_class = "hebrew-content"
    else:
        food_content = soup.find("div", id="food-en")
        content_class = "english-content"

    if not food_content:
        return doc

    # Get all categories
    categories = food_content.find_all("div", class_="category")

    for category in categories:
        # Skip quick navigation
        if category.find("div", class_="quick-nav"):
            continue

        # Category header
        category_header = category.find("div", class_="category-header")
        if category_header:
            cat_text = category_header.get_text(strip=True)
            # Remove emojis
            cat_text = re.sub(r"[^\w\s]", "", cat_text)
            if cat_text:
                heading = doc.add_heading(cat_text, level=1)
                heading_run = heading.runs[0]
                heading_run.font.size = Pt(16)
                heading_run.font.color.rgb = RGBColor(46, 125, 50)
                heading_run.bold = True

        # Add-ons box
        add_ons = category.find("div", class_="add-ons-box")
        if add_ons:
            add_ons_text = add_ons.get_text(strip=True)
            para = doc.add_paragraph(add_ons_text)
            para_run = para.runs[0]
            para_run.font.size = Pt(10)
            para_run.italic = True
            para_run.font.color.rgb = RGBColor(102, 102, 102)
            doc.add_paragraph()  # Empty line

        # Menu items
        menu_items = category.find_all("div", class_="menu-item")
        for item in menu_items:
            item_name = item.find("span", class_="item-name")
            item_price = item.find("span", class_="item-price")
            item_desc = item.find("div", class_="item-description")

            if item_name:
                name_text = item_name.get_text(strip=True)
                # Remove new badge text but keep the item name
                name_text = re.sub(r"חדש|NEW", "", name_text).strip()
                name_text = re.sub(r"\s+", " ", name_text)

                price_text = item_price.get_text(strip=True) if item_price else ""
                desc_text = item_desc.get_text(strip=True) if item_desc else ""

                # Item name and price
                para = doc.add_paragraph()
                name_run = para.add_run(name_text)
                name_run.bold = True
                name_run.font.size = Pt(12)

                if price_text:
                    para.add_run(" - " + price_text)
                    price_run = para.runs[-1]
                    price_run.font.color.rgb = RGBColor(46, 125, 50)
                    price_run.bold = True

                # Description
                if desc_text:
                    desc_para = doc.add_paragraph(desc_text)
                    desc_run = desc_para.runs[0]
                    desc_run.font.size = Pt(10)
                    desc_run.font.color.rgb = RGBColor(102, 102, 102)
                    desc_run.italic = True

                doc.add_paragraph()  # Empty line between items

    # Drinks section
    if language == "he":
        drinks_content = soup.find("div", id="drinks")
    else:
        drinks_content = soup.find("div", id="drinks-en")

    if drinks_content:
        doc.add_page_break()
        drinks_heading = doc.add_heading("משקאות" if language == "he" else "Drinks", 0)
        drinks_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        drinks_heading_run = drinks_heading.runs[0]
        drinks_heading_run.font.size = Pt(20)
        drinks_heading_run.font.color.rgb = RGBColor(46, 125, 50)
        drinks_heading_run.bold = True

        drink_categories = drinks_content.find_all("div", class_="category")
        for category in drink_categories:
            category_header = category.find("div", class_="category-header")
            if category_header:
                cat_text = category_header.get_text(strip=True)
                cat_text = re.sub(r"[^\w\s]", "", cat_text)
                if cat_text:
                    heading = doc.add_heading(cat_text, level=1)
                    heading_run = heading.runs[0]
                    heading_run.font.size = Pt(16)
                    heading_run.font.color.rgb = RGBColor(46, 125, 50)
                    heading_run.bold = True

            # Drink items
            drink_items = category.find_all("div", class_="drink-item")
            for item in drink_items:
                drink_name = item.find("div", class_="drink-name")
                drink_price = item.find("span", class_="item-price")
                drink_desc = item.find("div", class_="drink-description")

                if drink_name:
                    name_text = drink_name.get_text(strip=True)
                    price_text = drink_price.get_text(strip=True) if drink_price else ""
                    desc_text = drink_desc.get_text(strip=True) if drink_desc else ""

                    para = doc.add_paragraph()
                    name_run = para.add_run(name_text)
                    name_run.bold = True
                    name_run.font.size = Pt(12)

                    if price_text:
                        para.add_run(" - " + price_text)
                        price_run = para.runs[-1]
                        price_run.font.color.rgb = RGBColor(46, 125, 50)
                        price_run.bold = True

                    if desc_text:
                        desc_para = doc.add_paragraph(desc_text)
                        desc_run = desc_para.runs[0]
                        desc_run.font.size = Pt(10)
                        desc_run.font.color.rgb = RGBColor(102, 102, 102)
                        desc_run.italic = True

                    doc.add_paragraph()

    # Desserts section
    if language == "he":
        desserts_content = soup.find("div", id="desserts")
    else:
        desserts_content = soup.find("div", id="desserts-en")

    if desserts_content:
        doc.add_page_break()
        desserts_heading = doc.add_heading(
            "קינוחים" if language == "he" else "Desserts", 0
        )
        desserts_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desserts_heading_run = desserts_heading.runs[0]
        desserts_heading_run.font.size = Pt(20)
        desserts_heading_run.font.color.rgb = RGBColor(46, 125, 50)
        desserts_heading_run.bold = True

        dessert_items = desserts_content.find_all("div", class_="menu-item")
        for item in dessert_items:
            item_name = item.find("span", class_="item-name")
            item_price = item.find("span", class_="item-price")
            item_desc = item.find("div", class_="item-description")

            if item_name:
                name_text = item_name.get_text(strip=True)
                name_text = re.sub(r"חדש|NEW", "", name_text).strip()
                name_text = re.sub(r"\s+", " ", name_text)

                price_text = item_price.get_text(strip=True) if item_price else ""
                desc_text = item_desc.get_text(strip=True) if item_desc else ""

                para = doc.add_paragraph()
                name_run = para.add_run(name_text)
                name_run.bold = True
                name_run.font.size = Pt(12)

                if price_text:
                    para.add_run(" - " + price_text)
                    price_run = para.runs[-1]
                    price_run.font.color.rgb = RGBColor(46, 125, 50)
                    price_run.bold = True

                if desc_text:
                    desc_para = doc.add_paragraph(desc_text)
                    desc_run = desc_para.runs[0]
                    desc_run.font.size = Pt(10)
                    desc_run.font.color.rgb = RGBColor(102, 102, 102)
                    desc_run.italic = True

                doc.add_paragraph()

    # Footer info
    footer = soup.find("footer")
    if footer:
        doc.add_page_break()
        footer_heading = doc.add_heading(
            "צרו קשר" if language == "he" else "Contact", 0
        )
        footer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        phone = footer.find("a", class_="phone-number")
        if phone:
            doc.add_paragraph(phone.get_text(strip=True))

        hours = footer.find("div", class_="hours-section")
        if hours:
            doc.add_paragraph(hours.get_text(strip=True))

    return doc


# Create Hebrew document
print("Creating Hebrew menu document...")
hebrew_doc = create_document("he")
hebrew_doc.save("Thai_Garden_Menu_Hebrew.docx")
print("✓ Hebrew menu created: Thai_Garden_Menu_Hebrew.docx")

# Create English document
print("Creating English menu document...")
english_doc = create_document("en")
english_doc.save("Thai_Garden_Menu_English.docx")
print("✓ English menu created: Thai_Garden_Menu_English.docx")

print("\nDone! Both documents are ready for printing.")
